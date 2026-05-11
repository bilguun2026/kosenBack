import io
import re
import fitz  # PyMuPDF
from docx import Document
from django.http import JsonResponse
from django.views.decorators.http import require_POST


def is_bold_font(span):
    """Detect bold from flags OR font name (Medium, Bold, Black, Semibold)."""
    if span["flags"] & (1 << 4):  # bold flag
        return True
    font = span.get("font", "").lower()
    return any(w in font for w in ["bold", "medium", "black", "semibold", "heavy"])


def extract_pdf_text(file_obj):
    """
    PyMuPDF-based extraction that preserves headings, bold, lists, tables, and structure.
    """
    html_parts = []
    doc = fitz.open(stream=file_obj.read(), filetype="pdf")

    for page_num in range(len(doc)):
        page = doc[page_num]

        # --- Tables ---
        tabs = page.find_tables()
        table_rects = [t.bbox for t in tabs] if tabs else []

        for t in tabs:
            table_data = t.extract()
            if not table_data:
                continue
            html_parts.append(
                '<table style="border-collapse:collapse;width:100%;margin:12px 0;">'
            )
            for i, row in enumerate(table_data):
                html_parts.append("<tr>")
                for cell in row:
                    text = (cell or "").strip()
                    tag = "th" if i == 0 else "td"
                    style = 'style="border:1px solid #ccc;padding:8px 12px;text-align:left;'
                    if i == 0:
                        style += "background:#f0f4ff;font-weight:bold;"
                    style += '"'
                    html_parts.append(f"<{tag} {style}>{text}</{tag}>")
                html_parts.append("</tr>")
            html_parts.append("</table>")

        # --- Text blocks ---
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if "lines" not in block:
                continue

            # Skip text inside table areas
            bx0, by0, bx1, by1 = block["bbox"]
            in_table = False
            for tr in table_rects:
                tx0, ty0, tx1, ty1 = tr
                if bx0 >= tx0 - 2 and by0 >= ty0 - 2 and bx1 <= tx1 + 2 and by1 <= ty1 + 2:
                    in_table = True
                    break
            if in_table:
                continue

            for line in block["lines"]:
                spans = line["spans"]
                if not spans:
                    continue

                full_text = "".join(s["text"] for s in spans).strip()
                if not full_text:
                    continue

                # Get dominant font properties
                main_span = max(spans, key=lambda s: len(s["text"]))
                size = main_span["size"]
                bold = is_bold_font(main_span)

                # Skip page numbers
                if re.match(r"^\d+\s*/\s*\d+$", full_text):
                    continue

                # Determine HTML tag based on size + bold
                if size >= 13:
                    html_parts.append(f"<h1>{full_text}</h1>")
                elif size >= 11.5:
                    html_parts.append(f"<h2>{full_text}</h2>")
                elif size >= 10 and bold:
                    html_parts.append(f"<h3>{full_text}</h3>")
                elif full_text.startswith(("•", "·", "- ", "– ")):
                    bullet_text = full_text.lstrip("•·-– ").strip()
                    html_parts.append(f"<li>{bullet_text}</li>")
                elif bold:
                    # Build inline bold spans
                    inline = build_inline_html(spans)
                    html_parts.append(f"<p><strong>{inline}</strong></p>")
                else:
                    inline = build_inline_html(spans)
                    html_parts.append(f"<p>{inline}</p>")

    doc.close()

    # Wrap consecutive <li> in <ul>
    result = "\n".join(html_parts)
    result = re.sub(
        r"((?:<li>.*?</li>\n?)+)",
        lambda m: "<ul>\n" + m.group(0) + "</ul>\n",
        result,
    )
    return result


def build_inline_html(spans):
    """Build HTML with inline bold/italic from multiple spans."""
    parts = []
    for s in spans:
        t = s["text"]
        if not t:
            continue
        bold = is_bold_font(s)
        italic = bool(s["flags"] & (1 << 1))

        if bold and italic:
            t = f"<strong><em>{t}</em></strong>"
        elif bold:
            t = f"<strong>{t}</strong>"
        elif italic:
            t = f"<em>{t}</em>"
        parts.append(t)
    return "".join(parts)


def extract_docx_text(file_obj):
    doc = Document(file_obj)
    html_parts = []
    in_list = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            continue

        style = para.style.name.lower() if para.style else ""

        if "heading 1" in style:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h1>{text}</h1>")
        elif "heading 2" in style:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h2>{text}</h2>")
        elif "heading 3" in style:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h3>{text}</h3>")
        elif "list" in style or "bullet" in style:
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{text}</li>")
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            runs_html = ""
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold:
                    t = f"<strong>{t}</strong>"
                if run.italic:
                    t = f"<em>{t}</em>"
                if run.underline:
                    t = f"<u>{t}</u>"
                runs_html += t
            html_parts.append(f"<p>{runs_html or text}</p>")

    if in_list:
        html_parts.append("</ul>")

    for table in doc.tables:
        html_parts.append(
            '<table style="border-collapse:collapse;width:100%;margin:12px 0;">'
        )
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                style = 'style="border:1px solid #ccc;padding:8px 12px;text-align:left;'
                if i == 0:
                    style += "background:#f0f4ff;font-weight:bold;"
                style += '"'
                html_parts.append(f"<{tag} {style}>{cell.text.strip()}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    return "\n".join(html_parts)


@require_POST
def import_file_content(request):
    if not request.user.is_staff:
        return JsonResponse({"error": "Нэвтрэх шаардлагатай"}, status=403)

    uploaded = request.FILES.get("file")
    if not uploaded:
        return JsonResponse({"error": "Файл оруулна уу"}, status=400)

    name = uploaded.name.lower()
    file_bytes = io.BytesIO(uploaded.read())

    try:
        if name.endswith(".pdf"):
            html = extract_pdf_text(file_bytes)
        elif name.endswith((".docx",)):
            html = extract_docx_text(file_bytes)
        elif name.endswith((".doc",)):
            return JsonResponse(
                {"error": ".doc форматыг дэмжихгүй. .docx болгон хадгална уу."},
                status=400,
            )
        else:
            return JsonResponse(
                {"error": "PDF эсвэл DOCX файл оруулна уу"}, status=400
            )
    except Exception as e:
        return JsonResponse(
            {"error": f"Файл уншихад алдаа гарлаа: {str(e)}"}, status=500
        )

    if not html.strip():
        return JsonResponse({"error": "Файлаас текст олдсонгүй"}, status=400)

    return JsonResponse({"html": html})
